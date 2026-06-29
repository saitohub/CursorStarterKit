"""
build_docx.py — Markdown（note記事ドラフト）を Word (.docx) に変換する。

対応記法:
  - YAML frontmatter (title / tags)
  - 見出し H1 / H2 / H3
  - 箇条書き ("- ")
  - 引用 ("> ")
  - 太字 **text**, インライン `code`
  - コードブロック (``` ... ```)
  - 画像 ![alt](images/xxx.png)   ← 本文に埋め込み
  - 水平線 ---

画像ファイルが見つからない場合はプレースホルダを自動生成して埋め込む。

使い方:
  python build_docx.py --input path/to/article.md --output path/to/article.docx

依存:
  pip install python-docx Pillow
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.stderr.write(
        "[ERROR] python-docx が必要です。`pip install python-docx` を実行してください。\n"
    )
    sys.exit(1)

try:
    from PIL import Image, ImageDraw, ImageFont
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False

# 自前ヘルパ（同ディレクトリ）
sys.path.insert(0, str(Path(__file__).parent))
try:
    from _fontutil import get_japanese_font  # noqa: E402
    FONTUTIL_AVAILABLE = True
except Exception:
    FONTUTIL_AVAILABLE = False


FRONTMATTER_RE = re.compile(r"^---\s*\n(.*?)\n---\s*\n", re.DOTALL)
IMAGE_RE = re.compile(r"!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
INLINE_CODE_RE = re.compile(r"`([^`]+)`")


def parse_frontmatter(text: str) -> tuple[dict, str]:
    """frontmatter を dict で返し、残りの本文を返す。"""
    m = FRONTMATTER_RE.match(text)
    if not m:
        return {}, text

    raw = m.group(1)
    body = text[m.end():]
    meta: dict = {}

    for line in raw.splitlines():
        if ":" not in line:
            continue
        key, val = line.split(":", 1)
        key = key.strip()
        val = val.strip().strip('"').strip("'")
        if val.startswith("[") and val.endswith("]"):
            items = [v.strip().strip('"').strip("'") for v in val[1:-1].split(",")]
            meta[key] = [i for i in items if i]
        else:
            meta[key] = val

    return meta, body


def make_placeholder(path: Path, alt: str, size=(1280, 670)) -> bool:
    """画像が無い場合に灰色プレースホルダを作成。成功時 True。"""
    if not PIL_AVAILABLE:
        return False
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", size, color=(235, 238, 243))
    draw = ImageDraw.Draw(img)
    label = f"[PLACEHOLDER]\n{alt or path.name}"
    if FONTUTIL_AVAILABLE:
        font = get_japanese_font(size=40)
    else:
        try:
            font = ImageFont.truetype("arial.ttf", 40)
        except Exception:
            font = ImageFont.load_default()
    bbox = draw.multiline_textbbox((0, 0), label, font=font, spacing=8, align="center")
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.multiline_text(
        ((size[0] - w) / 2, (size[1] - h) / 2),
        label,
        font=font,
        fill=(120, 130, 150),
        spacing=8,
        align="center",
    )
    img.save(path, "PNG")
    return True


def add_inline(paragraph, text: str) -> None:
    """太字・インラインコードを解釈しつつ run を追加。"""
    idx = 0
    tokens: list[tuple[str, str]] = []
    for m in re.finditer(r"\*\*(.+?)\*\*|`([^`]+)`", text):
        if m.start() > idx:
            tokens.append(("plain", text[idx:m.start()]))
        if m.group(1) is not None:
            tokens.append(("bold", m.group(1)))
        else:
            tokens.append(("code", m.group(2)))
        idx = m.end()
    if idx < len(text):
        tokens.append(("plain", text[idx:]))

    for kind, val in tokens:
        run = paragraph.add_run(val)
        if kind == "bold":
            run.bold = True
        elif kind == "code":
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0x8E, 0x44, 0xAD)


def add_image(doc, img_path: Path, alt: str) -> None:
    """画像を挿入。無ければプレースホルダを生成。"""
    if not img_path.exists():
        created = make_placeholder(img_path, alt)
        if not created:
            para = doc.add_paragraph()
            run = para.add_run(f"[画像未配置: {img_path.name}] {alt}")
            run.italic = True
            run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
            return

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(str(img_path), width=Inches(5.8))

    if alt:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(alt)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)


def build_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    meta, body = parse_frontmatter(text)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Yu Gothic"
    style.font.size = Pt(11)

    title = meta.get("title") or md_path.stem
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    tags = meta.get("tags") or []
    if tags:
        tag_para = doc.add_paragraph()
        tag_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = tag_para.add_run("タグ: " + "  ".join(f"#{t}" for t in tags))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x77, 0xAA)

    doc.add_paragraph()

    lines = body.splitlines()
    in_code = False
    code_buffer: list[str] = []
    base_dir = md_path.parent

    for line in lines:
        stripped = line.rstrip()

        if stripped.startswith("```"):
            if in_code:
                code_para = doc.add_paragraph()
                code_para.paragraph_format.left_indent = Inches(0.3)
                run = code_para.add_run("\n".join(code_buffer))
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_buffer.append(line)
            continue

        if not stripped.strip():
            doc.add_paragraph()
            continue

        if stripped.strip() == "---":
            doc.add_paragraph("―" * 30)
            continue

        img_match = IMAGE_RE.match(stripped.strip())
        if img_match:
            rel = img_match.group("path")
            alt = img_match.group("alt")
            add_image(doc, (base_dir / rel).resolve(), alt)
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
            continue

        if stripped.lstrip().startswith("- "):
            content = stripped.lstrip()[2:]
            para = doc.add_paragraph(style="List Bullet")
            add_inline(para, content)
            continue

        if stripped.startswith("> "):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            run = para.add_run()
            add_inline(para, stripped[2:])
            for r in para.runs:
                r.italic = True
                r.font.color.rgb = RGBColor(0x60, 0x70, 0x80)
            continue

        para = doc.add_paragraph()
        add_inline(para, stripped)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> int:
    ap = argparse.ArgumentParser(description="Markdown→Word 変換（note記事用）")
    ap.add_argument("--input", required=True, help="入力 Markdown ファイル")
    ap.add_argument("--output", required=True, help="出力 .docx ファイル")
    args = ap.parse_args()

    md_path = Path(args.input).resolve()
    docx_path = Path(args.output).resolve()

    if not md_path.exists():
        sys.stderr.write(f"[ERROR] Markdown が見つかりません: {md_path}\n")
        return 1

    build_docx(md_path, docx_path)
    print(f"[OK] {docx_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
